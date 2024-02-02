import streamlit as st
from docx import Document
import io
import utils
import vertex

st.set_page_config(
    page_title=":sparkles: Contract liability predictor :sparkles:",
    page_icon=":robot:",
    layout="centered",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "# You will discover by yourself"
    }
)

#creating session states
utils.create_session_state()


st.title(":blue[PaLM 2 fine tuned] :blue[Rafa-Stephane-Julie] Contract Liability Predictor")



#with st.container():
    #st.write("Current Generator Settings: ")
    # if st.session_state['temperature'] or st.session_state['debug_mode'] or :
    #st.write ("Temperature: ",st.session_state['temperature']," \t \t Token limit: #",st.session_state['token_limit']
                #," \t \t Top-K: ",st.session_state['top_k']
                #," \t \t Top-P: ",st.session_state['top_p']
                #," \t \t Debug Model: ",st.session_state['debug_mode'])


    #prompt = st.text_area("Add your prompt: ",height = 100)
    
###trial julie

    #uploaded_files = st.file_uploader("Choose a docx contract file", accept_multiple_files=True)
    #for uploaded_file in uploaded_files:
        #bytes_data = uploaded_file.read()
        #st.write("filename:", uploaded_file.name)
        #st.write(bytes_data)
        


    # Function to read the content of a docx file
def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

    # Create the user interface to upload the file   
from docx import Document

# Function to read the content of a docx file and save it to a variable
def read_docx_and_save(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

st.title("Upload and Process DOCX File")

uploaded_file = st.file_uploader("Please upload your DOCX file here", type="docx")
if uploaded_file is not None:
    document_text = read_docx_and_save(uploaded_file)
    document_text="Given the following contract text, extract uncapped contract liabilities: "+ document_text
#####
    if document_text :
        #st.session_state['prompt'].append(prompt)
        st.markdown("<h3 style='text-align: center; color: blue;'>Generator Model Response</h3>", unsafe_allow_html=True)
        with st.spinner('DMZ Brain is working to generate, wait.....'):
            response = vertex.get_text_generation(prompt=document_text)
            #response = vertex.get_text_generation(prompt=prompt, temperature = st.session_state['temperature'],)
            st.session_state['response'].append(response)
            st.markdown(response)
